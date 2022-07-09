import csv
import datetime
import json
import os
import pathlib
import re
import subprocess
import uuid
from copy import deepcopy
from typing import List
import typer
import enum
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyPDF2 import PdfMerger
import time


class Result(enum.Enum):
    g = "g"
    e = "e"
    t = "t"


def create_list(paragraph, list_type):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), list_type) #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list
    pPr.append(numPr) #add number properties to paragraph


def insert_competences(competences, master):
    position: int = 0
    for index, paragraph in enumerate(master.paragraphs):
        if "inhalte" in paragraph.text.lower():
            position = index
    point = master.paragraphs[position + 2]
    for item in competences:
        p = point.insert_paragraph_before(style="List")
        p.style.font.size = Pt(14)
        p.paragraph_format.left_indent = Pt(30)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.line_spacing = 0.25
        r = p.add_run("•   " + item)
        r.font.name = "Constantia"
        r.font.size = Pt(14)
        r.font.bold = True
        create_list(p, "1")


def cleanup(files, callback, new_suffix: str = None):
    callback(1)
    for index, item in enumerate(files):
        # os.remove(str(item))
        if new_suffix:
            files[index] = item.with_suffix(new_suffix)
        callback(1)


def fill_template(doc, data):
    tokenRegex = re.compile(r'~@\[(?P<token>\w+)\]@~')
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            result = tokenRegex.findall(run.text)
            for token in result:
                if token in data.keys():
                    text = run.text.replace(f'~@[{token}]@~', data[token])
                    run.text = text
        print(paragraph.text)
    return doc


def main():
    stufen = {
        "g": "mit gutem Erfolg ",
        "e": "mit Erfolg ",
        "t": "",
    }
    dataset = []

    loading_start = time.time()
    with pathlib.Path("input.csv").open("r") as file:
        lc = len(file.readlines())
        file.seek(0)
        reader = csv.reader(file)
        with typer.progressbar(reader, length=lc, label="Laden        ", width=80, update_min_steps=10, show_eta=False,
                               empty_char=" ", show_pos=True) as bar:
            for row in bar:
                dataset.append({
                    "StudentFirstName": row[0],
                    "StudentLastName": row[1],
                    "StudentName": f'{row[0]} {row[1]}',
                    "StudentClass": row[2],
                    "StudentResult": stufen[row[3]],
                    "DocumentID": uuid.uuid1().hex
                })

    with pathlib.Path("config.json").open("r") as file:
        common = json.load(file)
        common["Date"] = datetime.date.today().strftime("%d.%m.%Y")
        common["TeacherGender"] = "" if common["TeacherGender"] == "m" else "in"
    with pathlib.Path("competences.txt").open("r") as file:
        competences = file.readlines()
    loading_end = time.time()

    files: List[pathlib.Path] = []
    generating_start = time.time()
    with typer.progressbar(dataset, label="Generieren   ", width=80, update_min_steps=5, show_pos=True, show_eta=False,
                           empty_char=" ") as bar:
        for unique in bar:
            doc = Document("Zertifikatvorlage_Technologie_G_2021_2022.docx")

            doc = fill_template(doc, unique)
            doc = fill_template(doc, common)
            insert_competences(competences, doc)

            output_dir = pathlib.Path("./tmp")
            if not output_dir.exists():
                output_dir.mkdir(parents=True)
            file = output_dir.joinpath(f"Cert_{unique['StudentFirstName']}_{unique['StudentLastName']}_{common['Term'].replace('/', '_')}.docx")
            doc.save(str(file))
            files.append(file)
    generating_end = time.time()

    converting_start = time.time()
    with typer.progressbar(files, label="Convertieren ", width=80, show_eta=True, show_pos=True, empty_char=" ") as bar:
        p = subprocess.Popen(f"libreoffice --headless --convert-to pdf --outdir {str(output_dir)} " +
                             " ".join([str(i) for i in files]),
                             shell=True,
                             stdout=subprocess.PIPE,
                             encoding="utf-8")
        while True:
            out = p.stdout.readline()
            if out == "" and p.poll() is not None:
                break
            if out.startswith("convert"):
                bar.update(1)
    converting_end = time.time()

    cleaning1_start = time.time()
    with typer.progressbar(length=len(files), label="Aufräumen    ", width=80, show_eta=False, show_pos=True,
                           empty_char=" ") as bar:
        cleanup(files, lambda x: bar.update(x), ".pdf")
    cleaning1_end = time.time()

    merging_start = time.time()
    merger = PdfMerger()
    with typer.progressbar(files,label="Zusammenfügen", width=80, show_pos=True, show_eta=False) as bar:
        for file in bar:
            merger.append(file)
    merger.write(str(output_dir.joinpath("Zertifikate.pdf")))
    merging_end = time.time()

    cleaning2_start = time.time()
    with typer.progressbar(files, label="Aufräumen    ", width=80, show_eta=False, show_pos=True,
                           empty_char=" ") as bar:
        cleanup(files, lambda x: bar.update(x))
    cleaning2_end = time.time()

    print(f'Laden:           {datetime.timedelta(seconds=loading_end - loading_start)}\tGenerieten:    {datetime.timedelta(seconds=generating_end - generating_start)}')
    print(f'Convertieren:    {datetime.timedelta(seconds=converting_end - converting_start)}\tZusammenfügen: {datetime.timedelta(seconds=merging_end - merging_start)}')
    print(f'Aufräumen:       {datetime.timedelta(seconds=(cleaning2_end - cleaning2_start) + (cleaning1_end - cleaning1_start))}')
    print(f'Gesammtlaufzeit: {datetime.timedelta(seconds=cleaning2_end - loading_start)}')


if __name__ == '__main__':
    typer.run(main)